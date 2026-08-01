#!/usr/bin/env python3
"""Render architecture Mermaid diagrams and export REVIEW_PIPELINE_ARCHITECTURE.docx."""

from __future__ import annotations

import os
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "docs" / "REVIEW_PIPELINE_ARCHITECTURE.md"
OUT = ROOT / "docs" / "REVIEW_PIPELINE_ARCHITECTURE.docx"
CHART_DIR = ROOT / "docs" / "architecture-charts"
MMDC = ROOT / "tools" / "mermaid" / "node_modules" / ".bin" / "mmdc"
NODE_BIN = ROOT / "tools" / "node" / "bin"
CHROME = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"

SLUGS = [
    "01-context",
    "02-components",
    "03-domain-er",
    "04-batch-state",
    "05-e2e-dataflow",
    "06-seq-match",
    "07-seq-defect",
    "08-deploy",
]

INIT = (
    "%%{init: {"
    "'flowchart': {"
    "  'nodeSpacing': 28, 'rankSpacing': 32, 'padding': 12, "
    "  'htmlLabels': true, 'wrappingWidth': 240"
    "}, "
    "'themeVariables': {"
    "  'fontSize': '18px', "
    "  'fontFamily': 'PingFang SC, Microsoft YaHei, sans-serif', "
    "  'lineColor': '#333333'"
    "}"
    "}}%%\n"
)

PORTRAIT_W, PORTRAIT_H = 6.5, 8.8
LANDSCAPE_W, LANDSCAPE_H = 9.8, 6.0
MIN_W = 5.5


def set_run_font(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), "PingFang SC")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    if color:
        run.font.color.rgb = RGBColor(*color)


def extract_mermaid(text: str) -> list[str]:
    blocks: list[str] = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        if lines[i].strip() == "```mermaid":
            i += 1
            body: list[str] = []
            while i < len(lines) and lines[i].strip() != "```":
                body.append(lines[i])
                i += 1
            blocks.append("\n".join(body).strip() + "\n")
        i += 1
    return blocks


def render_all(blocks: list[str]) -> list[Path]:
    if not MMDC.exists():
        raise SystemExit(f"mmdc not found: {MMDC}")
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    for p in CHART_DIR.glob("*"):
        p.unlink()

    env = os.environ.copy()
    env["PATH"] = f"{NODE_BIN}:{env.get('PATH', '')}"
    env["PUPPETEER_EXECUTABLE_PATH"] = CHROME

    paths: list[Path] = []
    for idx, src in enumerate(blocks):
        slug = SLUGS[idx] if idx < len(SLUGS) else f"diagram-{idx + 1:02d}"
        # erDiagram / stateDiagram / sequenceDiagram: skip flowchart-only init
        kind = src.lstrip().split(None, 1)[0]
        if kind.startswith("flowchart") or kind.startswith("graph"):
            body = INIT + src
        else:
            body = (
                "%%{init: {'themeVariables': {"
                "'fontSize': '16px', "
                "'fontFamily': 'PingFang SC, Microsoft YaHei, sans-serif'"
                "}}}%%\n"
                + src
            )
        mmd = CHART_DIR / f"{slug}.mmd"
        png = CHART_DIR / f"{slug}.png"
        mmd.write_text(body, encoding="utf-8")
        cmd = [
            str(MMDC),
            "-i",
            str(mmd),
            "-o",
            str(png),
            "-b",
            "white",
            "-s",
            "2.5",
            "-w",
            "1600",
        ]
        print("render", slug, flush=True)
        subprocess.run(cmd, check=True, env=env, cwd=str(ROOT / "tools" / "mermaid"))
        if not png.exists():
            raise SystemExit(f"missing {png}")
        paths.append(png)
        print(f"  ok {png.name} ({png.stat().st_size} bytes)")
    return paths


def choose_fit(png: Path) -> tuple[str, float, float]:
    with Image.open(png) as im:
        w_px, h_px = im.size
    aspect = h_px / max(w_px, 1)
    if aspect < 0.55:
        h = max(2.2, min(LANDSCAPE_H, LANDSCAPE_W * aspect))
        w = h / max(aspect, 1e-6)
        if w > LANDSCAPE_W:
            w, h = LANDSCAPE_W, LANDSCAPE_W * aspect
        return "landscape", w, h
    w, h = PORTRAIT_W, PORTRAIT_W * aspect
    if h <= PORTRAIT_H:
        return "portrait", w, h
    return "portrait", MIN_W, MIN_W * aspect


def set_section_orient(section, orient: str):
    if orient == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)


def add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    sizes = {1: 16, 2: 13, 3: 12}
    for run in h.runs:
        set_run_font(run, size=sizes.get(level, 12), bold=True, color=(0x1A, 0x1A, 0x1A))


def add_para(doc: Document, text: str, *, bold: bool = False, size: int = 10.5, space_after: int = 6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    # simple **bold** segments
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = p.add_run(part)
            set_run_font(run, size=size, bold=bold)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        set_run_font(cell.paragraphs[0].add_run(header), size=10, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i + 1].cells[c_i]
            cell.text = ""
            set_run_font(cell.paragraphs[0].add_run(val), size=9.5)
    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    header = [c.strip() for c in lines[start].strip().strip("|").split("|")]
    i = start + 1
    if i < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i].strip()):
        i += 1
    rows: list[list[str]] = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
        i += 1
    return header, rows, i


def build_docx(text: str, pngs: list[Path]) -> None:
    doc = Document()
    set_section_orient(doc.sections[0], "portrait")
    diagram_i = 0
    lines = text.splitlines()
    i = 0

    # Title from first H1
    while i < len(lines) and not lines[i].startswith("# "):
        i += 1
    if i < len(lines):
        title = lines[i][2:].strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(title), size=22, bold=True)
        i += 1

    while i < len(lines):
        line = lines[i]

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("> "):
            # blockquote: gather consecutive
            chunks: list[str] = []
            while i < len(lines) and lines[i].startswith("> "):
                chunks.append(lines[i][2:].rstrip())
                i += 1
            add_para(doc, " ".join(chunks), size=10, space_after=8)
            continue

        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 1)
            i += 1
            continue

        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2)
            i += 1
            continue

        if line.strip() == "```mermaid":
            i += 1
            while i < len(lines) and lines[i].strip() != "```":
                i += 1
            i += 1  # closing fence
            if diagram_i < len(pngs):
                png = pngs[diagram_i]
                diagram_i += 1
                orient, w, h = choose_fit(png)
                if orient == "landscape":
                    sec = doc.add_section()
                    set_section_orient(sec, "landscape")
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(png), width=Inches(w), height=Inches(h))
                if orient == "landscape":
                    sec2 = doc.add_section()
                    set_section_orient(sec2, "portrait")
            continue

        if line.strip().startswith("```"):
            # fenced non-mermaid (none expected) — skip
            i += 1
            while i < len(lines) and lines[i].strip() != "```":
                i += 1
            i += 1
            continue

        if line.strip().startswith("|") and "|" in line[1:]:
            headers, rows, i = parse_table(lines, i)
            add_table(doc, headers, rows)
            continue

        if re.match(r"^\d+\.\s+", line.strip()):
            text_item = re.sub(r"^\d+\.\s+", "", line.strip())
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            parts = re.split(r"(\*\*[^*]+\*\*)", text_item)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    set_run_font(run, size=10.5, bold=True)
                else:
                    run = p.add_run(part)
                    set_run_font(run, size=10.5)
            i += 1
            continue

        if line.strip().startswith("- "):
            text_item = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            parts = re.split(r"(\*\*[^*]+\*\*)", text_item)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    set_run_font(run, size=10.5, bold=True)
                else:
                    run = p.add_run(part)
                    set_run_font(run, size=10.5)
            i += 1
            continue

        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            add_para(doc, line.strip().strip("*"), size=9.5, space_after=8)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        add_para(doc, line.strip())
        i += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT} ({OUT.stat().st_size} bytes), diagrams={diagram_i}")


def main() -> None:
    text = MD.read_text(encoding="utf-8")
    blocks = extract_mermaid(text)
    if not blocks:
        raise SystemExit("no mermaid blocks")
    print(f"found {len(blocks)} diagrams")
    pngs = render_all(blocks)
    build_docx(text, pngs)


if __name__ == "__main__":
    main()
