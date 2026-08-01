#!/usr/bin/env python3
"""Build REVIEW_PIPELINE_FLOW.docx — prefer readable Chinese text over forcing tiny one-page charts."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "REVIEW_PIPELINE_FLOW.docx"
CHART_DIR = ROOT / "docs" / "flowcharts"
MANIFEST = CHART_DIR / "manifest.txt"

# Portrait usable; never crush Chinese text by shrinking too narrow.
PORTRAIT_W, PORTRAIT_H = 6.5, 9.0
LANDSCAPE_W, LANDSCAPE_H = 9.8, 6.2
MIN_READABLE_W = 5.5
MIN_READABLE_H = 2.4


def set_run_font(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), "PingFang SC")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    if color:
        run.font.color.rgb = RGBColor(*color)


def load_manifest() -> dict[str, Path]:
    out: dict[str, Path] = {}
    for line in MANIFEST.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        slug, _heading, name = line.split("\t", 2)
        out[slug] = CHART_DIR / name
    return out


def choose_fit(png: Path) -> tuple[str, float, float]:
    """Return (orient, width_in, height_in). Prefer readable Chinese text size."""
    with Image.open(png) as im:
        w_px, h_px = im.size
    aspect = h_px / max(w_px, 1)

    # Very wide → landscape, keep a minimum height so glyphs stay large enough
    if aspect < 0.55:
        h = max(MIN_READABLE_H, min(LANDSCAPE_H, LANDSCAPE_W * aspect))
        w = h / max(aspect, 1e-6)
        if w > LANDSCAPE_W:
            w = LANDSCAPE_W
            h = w * aspect
        return "landscape", w, h

    # Portrait: full width first
    w = PORTRAIT_W
    h = w * aspect
    if h <= PORTRAIT_H:
        return "portrait", w, h

    # Too tall: keep readable width (may overflow page — better than crush)
    return "portrait", MIN_READABLE_W, MIN_READABLE_W * aspect


def set_section_orient(section, orient: str):
    if orient == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)


def main() -> None:
    if not MANIFEST.exists():
        raise SystemExit(f"missing {MANIFEST}; run scripts/render_flowcharts.py first")

    by_slug = load_manifest()
    doc = Document()
    set_section_orient(doc.sections[0], "portrait")

    def page_break():
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

    def new_section(orient: str):
        """Add a new section with orientation (creates page break)."""
        new_sec = doc.add_section()
        set_section_orient(new_sec, orient)
        return new_sec

    def h1(text: str):
        heading = doc.add_heading(text, level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)
        for run in heading.runs:
            set_run_font(run, size=14, bold=True, color=(0x1A, 0x1A, 0x1A))

    def para(text: str, bold: bool = False, size: int = 10):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)

    def table(headers: list[str], rows: list[list[str]]):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        for i, header in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ""
            set_run_font(cell.paragraphs[0].add_run(header), size=10, bold=True)
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = t.rows[r_i + 1].cells[c_i]
                cell.text = ""
                set_run_font(cell.paragraphs[0].add_run(val), size=10)
        doc.add_paragraph()

    def figure(title: str, slug: str, note: str | None = None):
        png = by_slug.get(slug)
        if png is None or not png.exists():
            new_section("portrait")
            h1(title)
            para(f"[缺少流程图: {slug}]", bold=True)
            return
        orient, width, height = choose_fit(png)
        new_section(orient)
        h1(title)
        if note:
            para(note, size=10)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(png), width=Inches(width), height=Inches(height))

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run("成像专项评审流程图"), size=22, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        sub.add_run("大字号可读版 · 过高图自动横版 · 2026-07-30"),
        size=11,
        color=(0x55, 0x55, 0x55),
    )
    para(
        "说明：优先保证中文节点可读（不再为塞进竖版一页而压窄）。"
        "过高流程图自动改为横向页面；上传飞书请用本文件或 docs/flowcharts 下高清 PNG。",
        size=10,
    )

    figure("图 1 · 端到端总流程", "01-end-to-end")

    new_section("portrait")
    h1("图 2 · ① 任务加载方式（在线 vs 离线）")
    para("区别只在「任务包从哪来」；之后流程相同。")
    table(
        ["对比项", "在线加载", "离线加载"],
        [
            ["要不要现场上网", "要（连网 + 登录）", "不要"],
            ["任务从哪来", "刷新列表点选任务", "浏览器下 zip，手机选包"],
            ["关键操作", "刷新 → 选任务 → 加载", "选包 → 解压 → 加载"],
            ["场景进手机", "服务器拉取", "zip 解压"],
        ],
    )
    para("并排流程", bold=True)
    png = by_slug["02-load-mode-compare"]
    orient, w, h = choose_fit(png)
    # keep compare on this portrait section if possible
    if orient == "landscape":
        new_section("landscape")
        h1("图 2 · 并排流程（在线 vs 离线）")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(png), width=Inches(w), height=Inches(h))

    figure("图 2 · 怎么选加载方式", "02-load-mode-decide")
    figure(
        "图 2a · 在线加载",
        "02a-online-load",
        "路径：网络 → 登录 → 刷新 → 选任务 → 加载场景",
    )
    figure(
        "图 2b · 离线加载",
        "02b-offline-load",
        "路径：浏览器下 zip → 拷到手机 → 选包 → 解压 → 加载",
    )
    figure("图 3 · ② 标定设备角色", "03-device-role")
    figure("图 4 · ③ 按场景采集", "04-capture")

    new_section("portrait")
    h1("图 5 · ④ 上传通道决策")
    png = by_slug["05-upload-decide"]
    orient, w, h = choose_fit(png)
    if orient == "landscape":
        new_section("landscape")
        h1("图 5 · ④ 上传通道决策")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(png), width=Inches(w), height=Inches(h))
    if orient == "landscape":
        new_section("portrait")
    para("通道对照", bold=True)
    table(
        ["通道", "链路", "落点", "谁点上传"],
        [
            ["A", "拓展坞网口 → 服务器", "直达", "机端 App"],
            ["B1", "USB 共享网 → 服务器", "不落 PC", "机端 App"],
            ["B2", "ADB/panda → PC → 服务器", "先暂存再传", "PC 工具"],
        ],
    )

    figure("图 6 · 通道 A：拓展坞网口直连", "06-channel-a")
    figure("图 7 · 通道 B1：USB 共享网直传", "07-channel-b1")
    figure("图 8 · 通道 B2：拉取暂存再上传", "08-channel-b2")
    figure("图 9 · 三通道数据落点", "09-channel-sink")
    figure("图 10 · ⑤ 素材齐套检查", "10-readiness")
    figure(
        "图 11 · ⑥ 批量对比 → 建缺陷",
        "11-review-defect",
        "缺陷自动带入对应图片/视频；可先建一部分再提交。",
    )
    figure(
        "图 11b · ⑥c 服务器后台匹配",
        "11b-server-match",
        "数据上传完成后服务器后台跑；不阻塞评审。开发下载依赖匹配就绪。",
    )

    new_section("portrait")
    h1("图 12 · ⑦ 按模块指派并批量提交")
    png = by_slug["12-assign-submit"]
    orient, w, h = choose_fit(png)
    if orient == "landscape":
        new_section("landscape")
        h1("图 12 · ⑦ 按模块指派并批量提交")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(png), width=Inches(w), height=Inches(h))
    if orient == "landscape":
        new_section("portrait")
    table(
        ["步骤", "操作"],
        [
            ["选模块", "按成像模块归类缺陷"],
            ["选开发", "指定模块对应开发"],
            ["批量提交", "推到缺陷管理平台"],
        ],
    )

    figure(
        "图 13 · 开发取证：点图下载 log + dump",
        "13-dev-download",
        "点问题图下载「问题时段 log」+「图片对应 dump」。",
    )
    figure("图 14 · 异常回退", "14-exceptions")
    figure("图 15 · 角色分工", "15-roles")

    new_section("portrait")
    h1("附录")
    para("源文件：docs/REVIEW_PIPELINE_FLOW.md")
    para("高清 PNG（适合贴飞书）：docs/flowcharts/*.png")
    para("重新生成：python scripts/render_flowcharts.py && python scripts/export_review_flow_docx.py")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
